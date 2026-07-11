#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el Word KYC pre-llenado para Paul Velasquez Zavala
(socio y dueño formal de Qlick, persona física en RESICO 625)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ====== Datos del dueño (de la CSF) ======
PAUL = {
    "nombre_legal": "PAUL VELASQUEZ ZAVALA",
    "rfc": "VEZP810701CA8",
    "curp": "VEZP810701HBCLVL08",
    "fecha_nacimiento": "01/07/1981",
    "sexo": "Hombre",
    "estado_nacimiento": "Baja California",
    "nacionalidad": "Mexicana",
    "inicio_operaciones": "17/12/2004",
    "estatus": "ACTIVO",
    "direccion": "AV. CONSTITUCION S/N, PAREDONES, MEXICALI, BAJA CALIFORNIA, CP 21920",
    "cp": "21920",
    "municipio": "MEXICALI",
    "estado": "BAJA CALIFORNIA",
    "entre_calles": "Entre 2DA y BENITO JUAREZ",
    "regimen": "625 - Régimen Simplificado de Confianza (RESICO) - Persona Física",
    "regimen_codigo": "625",
    "actividades_publicidad": "5% Agencias de publicidad",
    "actividades_asalariado": "26% Asalariado",
    "actividades_principal": "Asalariado + Comercio al por menor (electród. y llantas)",
}

# ====== Datos del negocio (Qlick) ======
QLICK = {
    "nombre_legal_negocio": "Qlick",
    "sitio_web": "https://qlick.digital",
    "email_soporte": "hola@qlick.digital",
    "telefono_soporte": "(por confirmar con David)",
    "categoria_stripe": "Education & Training",
    "statement_descriptor": "QLICK.DIGITAL",
    "descripcion": (
        "Plataforma mexicana de educación online en marketing digital. "
        "Ofrece cursos grabados, workshops en vivo y material descargable "
        "para emprendedores y profesionales. Mercado: México y Latinoamérica "
        "hispanohablante. Sitio web: qlick.digital."
    ),
}

# ====== Estilos ======
COLOR_HEADER = RGBColor(0x00, 0x35, 0x54)
COLOR_WARN = RGBColor(0x99, 0x33, 0x33)
COLOR_OK = RGBColor(0x33, 0x77, 0x33)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_FILLED = "DCEBF7"   # azul claro para celdas pre-llenadas
COLOR_TODO = "FFE4B5"     # naranja claro para celdas pendientes
COLOR_HEADER_BG = "E7F0F4"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = COLOR_HEADER
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    return p


def add_warn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.font.color.rgb = COLOR_WARN
    return p


def add_ok(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("✓ " + text)
    run.font.color.rgb = COLOR_OK
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    return p


def add_meta(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    return p


def add_checkbox_paragraph(doc, label, checked=False):
    """Inserta un checkbox visual (no interactivo en LibreOffice, pero visible)."""
    mark = "☑" if checked else "☐"
    p = doc.add_paragraph()
    run = p.add_run(f"  {mark}  {label}")
    if checked:
        run.bold = True
    return p


def add_editable_cell(cell, value=None, todo=False):
    """Llena una celda. Si value es None, queda como placeholder."""
    if value is None:
        cell.text = "_____________ (llenar)"
        set_cell_bg(cell, COLOR_TODO)
    else:
        cell.text = value
        set_cell_bg(cell, COLOR_FILLED)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)


def make_table_with_sdt_checkboxes(doc, headers, rows, checkbox_col=None):
    """Crea tabla. Si checkbox_col es índice, agrega ☐/☑ en esa columna."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], COLOR_HEADER_BG)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_HEADER
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ====== Portada ======
    add_heading(doc, "Stripe KYC — Qlick (Persona Física · Paul Velasquez Zavala)", level=0)
    add_meta(doc, "Documento pre-llenado · Generado: 2026-07-10 · Régimen 625 RESICO confirmado")
    add_meta(doc, "Fuente: Constancia de Situación Fiscal del SAT (CSF) · Emisión: 10 de julio de 2026")

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Propósito: ").bold = True
    intro.add_run(
        "Completar la verificación de identidad (KYC) de Stripe México para que Qlick "
        "pueda recibir pagos con tarjeta en producción. La cuenta queda registrada "
        "a nombre de "
    )
    r = intro.add_run(PAUL["nombre_legal"])
    r.bold = True
    intro.add_run(
        ", persona física con actividad empresarial en régimen RESICO (código 625). "
        "El nombre comercial del negocio en Stripe será "
    )
    r = intro.add_run(QLICK["nombre_legal_negocio"])
    r.bold = True
    intro.add_run(".")

    p = doc.add_paragraph()
    p.add_run("Para quién es esta guía: ").bold = True
    p.add_run("la persona que operará el dashboard de Stripe (Paul o un delegado de su confianza).")

    p = doc.add_paragraph()
    p.add_run("Tiempo estimado: ").bold = True
    p.add_run("30–45 minutos si tiene los datos a mano; 1–2 horas si debe conseguirlos (CLABE, comprobante, fotos de INE).")

    # ====== Cómo usar ======
    add_heading(doc, "0. Cómo usar este documento", level=1)
    for i, t in enumerate([
        "Las celdas en AZUL CLARO contienen datos que ya están pre-llenados desde la CSF — verificar que coincidan carácter por carácter antes de enviar.",
        "Las celdas en NARANJA están pendientes — son datos que solo Paul puede proporcionar en el momento (CLABE, teléfono, email, etc.).",
        "Los ☐ son checkboxes visuales. Marcar con X o click si el editor lo soporta (Word 2010+ interactivo, versiones viejas marcar a mano).",
        "Toda la sección 3 (Datos personales) viene de la CSF. Si algo cambió (dirección, régimen), actualizar el CSF en el SAT antes de enviar a Stripe.",
    ], 1):
        doc.add_paragraph(f"  {i}. {t}", style="List Number")

    # ====== Sección 1: Tipo jurídico ======
    add_heading(doc, "1. Tipo jurídico del titular", level=1)
    add_checkbox_paragraph(doc, "Persona Física con Actividad Empresarial (Paul Velasquez Zavala)", checked=True)
    add_checkbox_paragraph(doc, "S.A. de C.V.", checked=False)
    add_checkbox_paragraph(doc, "S. de R.L. de C.V.", checked=False)
    add_note(doc, "La CSF confirma Persona Física (RFC de 13 caracteres: VEZP810701CA8). No hay acta constitutiva.")

    # ====== Sección 2: Régimen fiscal ======
    add_heading(doc, "2. Régimen fiscal declarado en el CSF", level=1)
    add_checkbox_paragraph(doc, "612 - Persona Física con Actividad Empresarial (régimen estándar)", checked=False)
    add_checkbox_paragraph(doc, "625 - Régimen Simplificado de Confianza (RESICO) — Persona Física  ←  QLICK", checked=True)
    add_checkbox_paragraph(doc, "601 - General de Ley Personas Morales (régimen estándar)", checked=False)
    add_checkbox_paragraph(doc, "626 - Régimen Simplificado de Confianza (RESICO) — Persona Moral", checked=False)
    add_checkbox_paragraph(doc, "603 - Personas Morales con Fines no Lucrativos (ONGs, A.C.)", checked=False)
    p = doc.add_paragraph()
    p.add_run("  ☐  Otro régimen (escribir código del SAT): ___________").bold = False
    add_note(doc, "La CSF lista Régimen Simplificado de Confianza desde 01/01/2022 (código SAT 625 para PF). "
                  "También aparece Régimen de Sueldos y Salarios (605), que aplica a su trabajo asalariado.")

    # ====== Sección 3: Datos personales (tabla pre-llenada) ======
    add_heading(doc, "3. Datos personales de Paul (pre-llenados desde CSF)", level=1)

    add_ok(doc, "Los siguientes datos salen DIRECTOS de la Constancia de Situación Fiscal. Verificar que coincidan carácter por carácter con la CSF y el INE.")
    add_warn(doc, "La dirección fiscal viene SIN número exterior (S/N) y SIN colonia. Si Stripe la rechaza, completar con número real o comprobante de domicilio alterno con dirección completa.")

    table = doc.add_table(rows=16, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Dato", "Valor"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], COLOR_HEADER_BG)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_HEADER

    rows = [
        ("1", "Nombre legal completo", PAUL["nombre_legal"]),
        ("2", "RFC (13 caracteres)", PAUL["rfc"]),
        ("3", "CURP", PAUL["curp"]),
        ("4", "Fecha de nacimiento (dd/mm/aaaa)", PAUL["fecha_nacimiento"]),
        ("5", "Sexo", PAUL["sexo"]),
        ("6", "Estado de nacimiento", PAUL["estado_nacimiento"]),
        ("7", "Nacionalidad", PAUL["nacionalidad"]),
        ("8", "Fecha de inicio de operaciones", PAUL["inicio_operaciones"]),
        ("9", "Estatus en el padrón", PAUL["estatus"]),
        ("10", "Régimen fiscal", PAUL["regimen"]),
        ("11", "Actividad económica declarada (la más relevante para Stripe)", PAUL["actividades_publicidad"]),
        ("12", "Domicilio fiscal (calle y número)", PAUL["direccion"].split(",")[0]),
        ("13", "Colonia / Localidad", "PAREDONES (sin colonia registrada en CSF — completar si Stripe lo pide)"),
        ("14", "CP / Municipio / Estado", f"CP {PAUL['cp']} · {PAUL['municipio']} · {PAUL['estado']}"),
        ("15", "Entre calles", PAUL["entre_calles"]),
    ]
    for ri, (num, field, val) in enumerate(rows):
        cells = table.rows[ri + 1].cells
        cells[0].text = num
        cells[1].text = field
        cells[2].text = val
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        set_cell_bg(cells[2], COLOR_FILLED)
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    add_note(doc, "Si la dirección fiscal es rechazada por Stripe, Paul debe presentar comprobante de domicilio alterno (luz, agua, estado de cuenta) con dirección completa incluyendo número exterior y colonia.")

    # ====== Sección 4: Datos pendientes de Paul (los que la CSF no tiene) ======
    add_heading(doc, "4. Datos que Paul debe proporcionar (no están en la CSF)", level=1)

    add_warn(doc, "Estos datos son obligatorios. No están en la CSF. Pedírselos a Paul directamente.")

    table = doc.add_table(rows=11, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Dato", "Valor"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], COLOR_HEADER_BG)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_HEADER

    pending_rows = [
        ("16", "Teléfono celular personal", "_____________"),
        ("17", "Email personal (el de uso diario, no el de Qlick)", "_____________"),
        ("18", "Estado civil", "_____________"),
        ("19", "Identificación oficial vigente (INE/IFE)", "_____________  (subir foto frente y vuelta, color, legible)"),
        ("20", "CLABE interbancaria (18 dígitos, cuenta a nombre de Paul Velasquez Zavala)", "_____________"),
        ("21", "Titular de la cuenta bancaria", "PAUL VELASQUEZ ZAVALA"),
        ("22", "Nombre del banco", "_____________  (BBVA, Banorte, Santander, HSBC, Scotiabank, Banregio — evitar neobanks)"),
        ("23", "Comprobante de domicilio (< 3 meses)", "_____________  (subir PDF o foto, con dirección completa)"),
        ("24", "Email de soporte del negocio (sugerido)", QLICK["email_soporte"]),
        ("25", "Teléfono de soporte del negocio", QLICK["telefono_soporte"]),
    ]
    for ri, (num, field, val) in enumerate(pending_rows):
        cells = table.rows[ri + 1].cells
        cells[0].text = num
        cells[1].text = field
        cells[2].text = val
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        # Si el valor es el placeholder, marcar en naranja. Si ya viene sugerido, en azul claro.
        if "_____________" in val:
            set_cell_bg(cells[2], COLOR_TODO)
        else:
            set_cell_bg(cells[2], COLOR_FILLED)
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    # ====== Sección 5: Datos del negocio (semi pre-llenados) ======
    add_heading(doc, "5. Datos del negocio en Stripe", level=1)

    table = doc.add_table(rows=7, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Dato", "Valor"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], COLOR_HEADER_BG)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_HEADER

    business_rows = [
        ("26", "Nombre del negocio (Business name)", QLICK["nombre_legal_negocio"]),
        ("27", "Categoría del negocio (Business category)", QLICK["categoria_stripe"]),
        ("28", "Descripción del negocio (100–300 palabras)", QLICK["descripcion"]),
        ("29", "Sitio web", QLICK["sitio_web"]),
        ("30", "Statement descriptor (aparece en estado de cuenta del cliente)", QLICK["statement_descriptor"]),
        ("31", "Email de soporte", QLICK["email_soporte"]),
    ]
    for ri, (num, field, val) in enumerate(business_rows):
        cells = table.rows[ri + 1].cells
        cells[0].text = num
        cells[1].text = field
        cells[2].text = val
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        set_cell_bg(cells[2], COLOR_FILLED)
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    # ====== Sección 6: Pasos del KYC ======
    add_heading(doc, "6. Pasos en el dashboard de Stripe", level=1)

    steps = [
        ("Paso 1 — Iniciar sesión", "https://dashboard.stripe.com/login — con las credenciales que David creó."),
        ("Paso 2 — Verificar modo Live", "Arriba a la derecha debe decir \"Live data\". Si dice \"Test data\", cambiar con el toggle. Stripe en modo test no procesa pagos reales."),
        ("Paso 3 — Iniciar KYC", "https://dashboard.stripe.com/account/onboarding — click en \"Add information to start accepting live payments\"."),
        ("Paso 4 — Tipo de cuenta", "Elegir \"I'm setting up a business\" (empresa, NO individual — si elige individual tendrá que reiniciar)."),
        ("Paso 5 — País y tipo jurídico", "País: Mexico. Tipo jurídico: Sole proprietorship (Persona Física con Actividad Empresarial)."),
        ("Paso 6 — Datos del titular", "Llenar con los puntos #1–#15 de la sección 3. Verificar que el nombre coincida carácter por carácter con el INE."),
        ("Paso 7 — Datos del negocio", "Business name: Qlick. Categoría: Education & Training. Descripción: la del punto #28. Sitio web: https://qlick.digital. Email y teléfono: los del punto #24 y #25."),
        ("Paso 8 — Statement descriptor", "https://dashboard.stripe.com/settings/public — escribir QLICK.DIGITAL (punto #30). Aparece en el estado de cuenta del cliente."),
        ("Paso 9 — Cuenta bancaria", "Payouts. País: Mexico. Moneda: MXN. Account number: CLABE del punto #20 (18 dígitos). Account holder: PAUL VELASQUEZ ZAVALA (debe coincidir carácter por carácter con la cuenta)."),
        ("Paso 10 — Subir documentos", "INE de Paul (foto frente y vuelta, color, legible). Comprobante de domicilio (< 3 meses). CSF si Stripe la pide. Buena luz, plano sobre mesa."),
        ("Paso 11 — Microdepósitos", "Stripe transfiere $1–$5 MXN a la cuenta de Paul. Tarda 1–3 días hábiles. Confirmar los montos en el dashboard cuando lleguen."),
        ("Paso 12 — Activar métodos de pago", "Cards (obligatorio), OXXO (recomendado), SPEI / customer balance (opcional). URL: https://dashboard.stripe.com/settings/payment_methods."),
        ("Paso 13 — Esperar activación", "https://dashboard.stripe.com/account — sección \"Account status\". 🟡 Restricted = falta algo. 🟢 Complete = listo. Avisar a David."),
    ]
    table = doc.add_table(rows=1 + len(steps), cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["✓", "Paso", "Descripción"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], COLOR_HEADER_BG)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_HEADER
    for i, (title, body) in enumerate(steps):
        cells = table.rows[i + 1].cells
        cells[0].text = "☐"
        cells[1].text = title
        cells[2].text = body
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        for p in cells[1].paragraphs:
            for r in p.runs:
                r.bold = True

    # ====== Sección 7: Gotchas ======
    add_heading(doc, "7. Gotchas importantes", level=1)
    add_warn(doc, "Nombre carácter por carácter. \"VELASQUEZ\" vs \"VELASQUES\" (sin Z) causa rechazo. El nombre en Stripe debe coincidir EXACTO con el INE y la CSF.")
    add_warn(doc, "Régimen 625 RESICO, NO 612. La CSF NO muestra el régimen 612. La actividad publicitaria de Paul está declarada bajo RESICO 625, no bajo 612. Si el operador pone 612 por error, el KYC se complica.")
    add_warn(doc, "Actividad preponderante. Las actividades con mayor porcentaje en la CSF son asalariado (26%), comercio de electrodomésticos (20%) y llantas (20%). \"Agencias de publicidad\" solo aparece al 5%. Si Stripe pregunta actividad principal, decir la realidad: asalariado + comercio, con la actividad publicitaria como secundaria. NO inventar que es 100% cursos de marketing.")
    add_warn(doc, "Dirección fiscal incompleta. La CSF no tiene número exterior (S/N) ni colonia. Si Stripe rechaza, subir comprobante de domicilio alterno (luz, agua, estado de cuenta) con dirección completa real de Paul.")
    add_warn(doc, "Cuenta bancaria a nombre de Paul. El titular de la CLABE debe ser PAUL VELASQUEZ ZAVALA, no Qlick, no David, no una empresa. Si Paul no tiene cuenta a su nombre, aperturar una antes de continuar.")
    add_warn(doc, "No neobanks. Nu, HeyBanco, Spin y Mercado Pago wallet a veces no funcionan con Stripe payouts. Si la cuenta de Paul es en uno de esos, abrir una en BBVA, Banorte o Santander tradicional.")
    add_warn(doc, "CLABE 18 dígitos, no 16. La CLABE interbancaria son 18 dígitos. El número de tarjeta son 16. No confundir.")
    add_warn(doc, "Una vez confirmado, no se puede editar. País, nombre legal, RFC y tipo jurídico no son editables. Verificar tres veces antes de avanzar.")

    # ====== Sección 8: Avisar a David ======
    add_heading(doc, "8. Qué avisarle a David cuando termine", level=1)
    for b in [
        "Screenshot del estado verde en https://dashboard.stripe.com/account",
        "Confirmación de que puede generar keys live en https://dashboard.stripe.com/apikeys (toggle \"Live data\")",
        "Métodos de pago activos: Cards (obligatorio), OXXO (recomendado), SPEI (opcional)",
        "URL del webhook endpoint que David va a registrar: https://www.qlick.digital/api/webhooks/stripe",
        "Eventos requeridos: checkout.session.completed, checkout.session.async_payment_succeeded, checkout.session.async_payment_failed, checkout.session.expired, charge.refunded",
    ]:
        doc.add_paragraph(b, style="List Number")

    # ====== Sección 9: Glosario ======
    add_heading(doc, "9. Glosario mínimo", level=1)
    glossary = [
        ("KYC", "Know Your Customer — verificación obligatoria de identidad por ley."),
        ("RFC", "Registro Federal de Contribuyentes (México). PF = 13 caracteres. Empresas = 12."),
        ("CURP", "Clave Única de Registro de Población. 18 caracteres. Identidad civil mexicana."),
        ("CSF", "Constancia de Situación Fiscal del SAT. Documento oficial con RFC, domicilio, régimen fiscal y obligaciones."),
        ("RESICO", "Régimen Simplificado de Confianza. Régimen fiscal opcional del SAT. Códigos: 625 (PF) o 626 (PM). Paul/Qlick está en 625."),
        ("Owner", "Persona con participación o control significativo. En PF con AE, el dueño es el titular."),
        ("Representante legal", "Persona autorizada para actuar en nombre del titular. En PF con AE, es el mismo Paul."),
        ("CLABE", "Clave Bancaria Estandarizada. 18 dígitos. La usa Stripe para depositar las ventas."),
        ("Statement descriptor", "Texto corto (5–22 chars) que aparece en el estado de cuenta del cliente. Para Qlick: QLICK.DIGITAL."),
        ("Microdepósito", "Transferencia pequeña ($1–$5 MXN) que Stripe hace para verificar propiedad de la cuenta bancaria."),
        ("Sole proprietorship", "Término en inglés de Stripe para Persona Física con Actividad Empresarial."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        p.add_run(term + ". ").bold = True
        p.add_run(defn)

    # ====== Sección 10: Confirmación a David ======
    add_heading(doc, "10. Confirmación a David", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Estado"
    for c in hdr:
        set_cell_bg(c, COLOR_HEADER_BG)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_HEADER

    final_rows = [
        ("KYC completado (estado 🟢 en dashboard)", "☐ Sí  ☐ Pendiente  ☐ Rechazado (especificar causa)"),
        ("Cuenta bancaria verificada (microdepósitos confirmados)", "☐ Sí  ☐ Pendiente"),
        ("Keys live generadas y entregadas a David", "☐ Sí  ☐ Pendiente"),
        ("Métodos de pago activos (Cards, OXXO, SPEI)", "☐ Cards  ☐ OXXO  ☐ SPEI"),
        ("Notas del operador (problemas, fricciones, rechazos):", "_____________________________________________"),
    ]
    for ri, (item, status) in enumerate(final_rows):
        cells = table.rows[ri + 1].cells
        cells[0].text = item
        cells[1].text = status
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    out = r"C:\Users\User\Documents\Click\docs\STRIPE_KYC_QLICK_PAUL_PRELLENADO.docx"
    doc.save(out)
    print(f"OK {out}")


if __name__ == "__main__":
    build()
