#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera UN solo Word con checkboxes interactivos + tabla PF vs Moral + campos de texto."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import uuid

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

DARK_BLUE = RGBColor(0x00, 0x35, 0x54)
RED = RGBColor(0x99, 0x33, 0x33)
GREEN = RGBColor(0x33, 0x77, 0x33)
GRAY = RGBColor(0x55, 0x55, 0x55)


def sdt_checkbox(checked=False, label_after=" "):
    """Returns an SDT checkbox element + an optional text node label."""
    state = "1" if checked else "0"
    xml = f'''
    <w:sdt xmlns:w="{W_NS}" xmlns:w14="{W14_NS}">
      <w:sdtPr>
        <w:rPr><w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic"/></w:rPr>
        <w:id w:val="{abs(hash(uuid.uuid4())) % 2147483647}"/>
        <w14:checkbox>
          <w14:checked w14:val="{state}"/>
          <w14:checkedState w14:val="2612" w14:font="MS Gothic"/>
          <w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/>
        </w14:checkbox>
      </w:sdtPr>
      <w:sdtContent>
        <w:r>
          <w:rPr><w:rFonts w:ascii="MS Gothic" w:eastAsia="MS Gothic" w:hAnsi="MS Gothic"/></w:rPr>
          <w:t>{'☒' if checked else '☐'}</w:t>
        </w:r>
      </w:sdtContent>
    </w:sdt>
    '''
    return etree.fromstring(xml)


def sdt_text(placeholder="Click para escribir..."):
    """Returns an SDT plain-text content control."""
    # Escape placeholder to avoid XML injection on '&' or '<' etc.
    safe = (placeholder
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))
    xml = (
        f'<w:sdt xmlns:w="{W_NS}">'
        f'<w:sdtPr>'
        f'<w:id w:val="{abs(hash(uuid.uuid4())) % 2147483647}"/>'
        f'<w:placeholder><w:docPart w:val="DefaultPlaceholder_-1854013440"/></w:placeholder>'
        f'<w:showingPlcHdr/>'
        f'<w:text/>'
        f'</w:sdtPr>'
        f'<w:sdtContent>'
        f'<w:r>'
        f'<w:rPr><w:color w:val="888888"/><w:i/></w:rPr>'
        f'<w:t xml:space="preserve">{safe}</w:t>'
        f'</w:r>'
        f'</w:sdtContent>'
        f'</w:sdt>'
    )
    return etree.fromstring(xml)


def add_checkbox_line(doc, label, checked=False, bold=False, size=None):
    """Adds a paragraph: [☐] label"""
    p = doc.add_paragraph()
    if size:
        p.paragraph_format.space_after = Pt(size)
    p._p.append(sdt_checkbox(checked=checked))
    run = p.add_run("  " + label)
    if bold:
        run.bold = True


def add_text_field(doc, placeholder="Click para escribir aquí"):
    """Adds a paragraph with a single text SDT."""
    p = doc.add_paragraph()
    p._p.append(sdt_text(placeholder))


def add_inline_text_field(paragraph, placeholder="Click para escribir"):
    """Appends a text field to an existing paragraph."""
    paragraph._p.append(sdt_text(placeholder))


def heading(doc, text, level=1, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(11)
    return p


def note(doc, text, color=RED, bold=False):
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.font.color.rgb = color
    if bold:
        run.bold = True


def ok(doc, text):
    note(doc, text, color=GREEN)


def meta(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def shade_header(table, hex_color="E7F0F4"):
    for cell in table.rows[0].cells:
        set_cell_bg(cell, hex_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = DARK_BLUE


def normal_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    shade_header(table)
    return table


def datos_table(doc):
    """Tabla de datos lado a lado: cada fila es 'Campo | PF | Moral | Notas' donde PF y Moral son text fields SDT."""
    headers = ["#", "Campo", "Persona Física (si aplica)", "Persona Moral (si aplica)", "Notas / Tips"]
    table = doc.add_table(rows=1 + 14, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    shade_header(table)

    # Helper to add an SDT to a cell's first paragraph
    def cell_with_sdt(cell, placeholder):
        # Clear cell text first
        cell.text = ""
        p = cell.paragraphs[0]
        sdt = sdt_text(placeholder)
        p._p.append(sdt)

    rows = [
        ("1", "Tipo jurídico exacto",
            "Persona Física con Actividad Empresarial",
            "S.A. de C.V. o S. de R.L. de C.V.",
            "Confirmar con acta constitutiva o alta de SAT"),
        ("2", "Nombre legal / Razón social",
            "Tu nombre completo tal como aparece en INE",
            "Razón social completa de la empresa",
            "PF: carácter por carácter vs INE. Moral: carácter por carácter vs CSF."),
        ("3", "RFC",
            "13 caracteres (persona física)",
            "12 caracteres (empresa)",
            "Sacar de CSF del SAT. Verificar que el RFC esté bien escrito."),
        ("4", "Régimen fiscal (CSF)",
            "612 PF Actividad Empresarial · 625 RESICO PF",
            "601 General Ley Moral · 626 RESICO Moral",
            "Qlick está en RESICO. PF → 625. Moral → 626. Confirmar en CSF."),
        ("5", "Domicilio fiscal",
            "Tu domicilio personal",
            "Domicilio fiscal de la empresa (del CSF)",
            "Calle + número + colonia + CP + municipio + estado + país"),
        ("6", "Fecha de constitución / nacimiento",
            "Tu fecha de nacimiento",
            "Fecha de constitución de la empresa",
            "dd/mm/aaaa"),
        ("7", "CURP",
            "Tu CURP (de INE)",
            "Solo si eres el representante legal → tu CURP",
            "Solo Moral lo pide para representante"),
        ("8", "INE / identificación",
            "Tu INE vigente",
            "INE del representante legal",
            "Frente y vuelta, color, legible"),
        ("9", "Comprobante de domicilio",
            "< 3 meses (estado de cuenta, CFE, Telmex)",
            "< 3 meses",
            "Solo Moral pide adicional comprobante de la empresa"),
        ("10", "CLABE interbancaria (18 dígitos)",
            "De tu cuenta personal",
            "De la cuenta de la empresa",
            "App del banco, sección datos para transferencia. 18 dígitos, no 16."),
        ("11", "Titular de la cuenta bancaria",
            "Tu nombre",
            "Razón social de la empresa",
            "Moral: la cuenta debe estar a nombre de la empresa, NO tuya"),
        ("12", "Banco",
            "BBVA / Banorte / Santander / HSBC / Scotiabank / Banregio",
            "Mismos bancos tradicionales",
            "Evitar neobanks: Nu, HeyBanco, Spin, Mercado Pago wallet"),
        ("13", "Statement descriptor (5–22 chars)",
            "Texto que aparece en estado de cuenta del cliente",
            "Texto que aparece en estado de cuenta del cliente",
            "Sugerencia: QLICK.DIGITAL o QLICK CURSOS"),
        ("14", "Sitio web / soporte / teléfono",
            "https://qlick.digital · hola@qlick.digital",
            "Mismos valores (públicos)",
            "David confirma los reales"),
    ]

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ci == 2 or ci == 3:
                # Persona Física o Persona Moral column → place a text field
                cell_with_sdt(cell, val if val else "Tocar para llenar")
            else:
                cell.text = val
    return table


def pasos_table(doc):
    """Tabla de pasos. Cada fila tiene un checkbox en la columna 'Hecho'."""
    headers = ["#", "Paso del KYC en Stripe", "URL", "Aplica a", "Hecho"]
    table = doc.add_table(rows=1 + 13, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    shade_header(table)

    data = [
        ("1", "Iniciar sesión en Stripe",
            "https://dashboard.stripe.com/login",
            "Ambos",
            "☐"),
        ("2", "Verificar que está en modo Live",
            "dashboard.stripe.com",
            "Ambos",
            "☐"),
        ("3", "Iniciar el KYC",
            "https://dashboard.stripe.com/account/onboarding",
            "Ambos",
            "☐"),
        ("4", "Tipo de cuenta: business",
            "(dentro del flujo)",
            "Ambos",
            "☐"),
        ("5", "País: Mexico + tipo jurídico",
            "(dentro del flujo)",
            "Ambos",
            "☐"),
        ("6", "Datos personales o de empresa",
            "(dentro del flujo)",
            "Ambos — campos distintos",
            "☐"),
        ("7", "Datos del negocio (categoría, descripción, web)",
            "(dentro del flujo)",
            "Ambos",
            "☐"),
        ("8", "Representante legal (solo Moral)",
            "(dentro del flujo)",
            "Solo Moral",
            "☐"),
        ("9", "Owners adicionales",
            "(dentro del flujo)",
            "Solo Moral si hay >25% / >15%",
            "☐"),
        ("10", "Statement descriptor",
            "https://dashboard.stripe.com/settings/public",
            "Ambos",
            "☐"),
        ("11", "Cuenta bancaria (CLABE)",
            "(dentro del flujo)",
            "Ambos — titular distinto",
            "☐"),
        ("12", "Subir documentos (INE, CSF)",
            "(dentro del flujo)",
            "Ambos — Moral sube más",
            "☐"),
        ("13", "Esperar activación",
            "https://dashboard.stripe.com/account",
            "Ambos",
            "☐"),
    ]

    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ci == 4:
                # Hecho column → SDT checkbox
                cell.text = ""
                p = cell.paragraphs[0]
                p._p.append(sdt_checkbox())
            else:
                cell.text = val
    return table


def main():
    doc = Document()
    # Document margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ----- Encabezado -----
    heading(doc, "Stripe KYC Qlick — Formulario único auto-seleccionable", level=0)
    meta(doc, "Qlick Marketing Digital · Última actualización: 2026-07-09 (incluye Régimen Simplificado de Confianza / RESICO)")
    meta(doc, "Pensado para que UN solo documento sirva tanto si Qlick es Persona Física como Persona Moral.")

    # ----- Instrucciones -----
    heading(doc, "Cómo usar este documento", level=1)
    instructions = [
        "Marque con qué tipo jurídico está constituida Qlick en la SECCIÓN 1.",
        "Marque el régimen fiscal del CSF en la SECCIÓN 2.",
        "Llene SOLO los datos que aplican a su tipo en la SECCIÓN 3 (las columnas están separadas).",
        "Vaya marcando los pasos completados en la SECCIÓN 4 a medida que avanza.",
        "Cuando la cuenta esté 🟢 Complete en Stripe, llene la SECCIÓN 5 y avísele a David.",
    ]
    for i, t in enumerate(instructions):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. ").bold = True
        p.add_run(t)

    # ----- Sección 1: Tipo jurídico -----
    heading(doc, "Sección 1 · Tipo jurídico de Qlick", level=1)
    p = doc.add_paragraph()
    p.add_run("Marque con qué tipo está constituida Qlick ").bold = True
    p.add_run("(marque solo uno):")

    add_checkbox_line(doc, "Persona Física con Actividad Empresarial")
    add_checkbox_line(doc, "Persona Moral — S.A. de C.V.", checked=False)
    add_checkbox_line(doc, "Persona Moral — S. de R.L. de C.V.", checked=False)

    note(doc,
        "Si marca Persona Física, ignore las columnas y filas de Persona Moral en la SECCIÓN 3. "
        "Si marca Persona Moral, ignore la columna de Persona Física. "
        "Si ninguna opción aplica exactamente, frene y avísele a David antes de avanzar.")

    # ----- Sección 2: Régimen fiscal -----
    heading(doc, "Sección 2 · Régimen fiscal declarado en el CSF del SAT", level=1)
    p = doc.add_paragraph()
    p.add_run("Marque el régimen fiscal que aparece en el CSF (Constancia de Situación Fiscal) de Qlick ").bold = True
    p.add_run("(marque solo uno):")

    add_checkbox_line(doc, "612 — Persona Física con Actividad Empresarial (régimen estándar)")
    add_checkbox_line(doc, "625 — Régimen Simplificado de Confianza (RESICO), persona física   ← QLICK si es PF con AE", bold=True)
    add_checkbox_line(doc, "601 — General de Ley Personas Morales (régimen estándar)")
    add_checkbox_line(doc, "626 — Régimen Simplificado de Confianza (RESICO), persona moral   ← QLICK si es Moral", bold=True)
    add_checkbox_line(doc, "603 — Personas Morales con Fines no Lucrativos (ONGs, A.C.)")

    p = doc.add_paragraph()
    p.add_run("Otro régimen (escribir código del SAT): ")
    add_inline_text_field(p, "Click para escribir (ej. 607, 610…)")

    ok(doc,
        "Qlick está dada de alta en RESICO según David. "
        "Si su CSF dice 625 está como persona física, si dice 626 está como persona moral. "
        "Si el CSF trae un código distinto, frene y avísele a David antes de avanzar en el KYC.")

    # ----- Sección 3: Datos a llenar -----
    heading(doc, "Sección 3 · Datos a llenar en Stripe", level=1)
    p = doc.add_paragraph()
    p.add_run("Toca cada celda azul para escribir. ").bold = True
    p.add_run("Llene SOLO los datos que aplican según su selección de la Sección 1:")

    datos_table(doc)

    note(doc,
        "Recordatorios. CLABE son 18 dígitos (los números de tarjeta son 16 y NO sirven). "
        "Si Qlick es Persona Moral, la CLABE y la cuenta bancaria deben estar a nombre de la empresa, NO tuya. "
        "El nombre legal debe coincidir carácter por carácter con el documento (INE para PF, CSF para Moral). "
        "Si Qlick es Persona Moral y solo vos sos dueño, indicar \"No\" cuando pregunte por owners adicionales "
        "o agregar a David con 100% de participación.")

    # ----- Sección 4: Pasos del KYC -----
    heading(doc, "Sección 4 · Pasos del KYC en Stripe", level=1)
    p = doc.add_paragraph()
    p.add_run("Marque cada paso cuando lo complete. ").bold = True
    p.add_run("Si un paso no aplica (ej. owners en PF), márquelo también con palomita y anote \"no aplica\":")

    pasos_table(doc)

    note(doc,
        "Si algún paso se queda en 🟡 Restricted, revise qué falta en la sección \"Account status\" de Stripe. "
        "Stripe transfiere $1–$5 MXN a tu cuenta para verificar propiedad (microdepósitos). "
        "Tarda 1–3 días hábiles, el proceso se pausa hasta confirmarlos.")

    # ----- Sección 5: Confirmación a David -----
    heading(doc, "Sección 5 · Cuando todo esté 🟢 — Avisar a David", level=1)
    p = doc.add_paragraph("Cuando el estado esté 🟢 Complete, envíale a David:")

    add_checkbox_line(doc, "Screenshot del estado verde en https://dashboard.stripe.com/account")
    add_checkbox_line(doc, "Confirmación de que se pueden generar keys live en https://dashboard.stripe.com/apikeys (toggle \"Live data\")")
    add_checkbox_line(doc, "Confirmación de que están activos: Cards (obligatorio), OXXO (recomendado), SPEI (opcional)")
    add_checkbox_line(doc, "URL del webhook endpoint: https://www.qlick.digital/api/webhooks/stripe — eventos: checkout.session.completed, checkout.session.async_payment_succeeded, checkout.session.async_payment_failed, checkout.session.expired, charge.refunded")
    add_checkbox_line(doc, "Marcar tipo jurídico final (Sección 1) y régimen fiscal final (Sección 2) para que David lo documente")

    # ----- Glosario -----
    heading(doc, "Glosario", level=1)
    glossary = [
        ("KYC", "\"Know Your Customer\" — verificación obligatoria de identidad por ley."),
        ("RFC", "Registro Federal de Contribuyentes. Personas físicas = 13 chars. Empresas = 12."),
        ("CSF", "Constancia de Situación Fiscal del SAT. PDF oficial con RFC, razón social, domicilio fiscal y régimen fiscal."),
        ("RESICO", "Régimen Simplificado de Confianza. Régimen fiscal opcional del SAT para PF y morales con ingresos menores a ciertos topes. Códigos: 626 (moral) o 625 (PF con AE)."),
        ("S.A. de C.V.", "Sociedad Anónima de Capital Variable. Tipo jurídico común para empresas en México."),
        ("S. de R.L. de C.V.", "Sociedad de Responsabilidad Limitada de Capital Variable."),
        ("CLABE", "Clave Bancaria Estandarizada. 18 dígitos. La usa Stripe para depositar las ventas."),
        ("Statement descriptor", "Texto corto que aparece en el estado de cuenta del cliente cuando le cobrás."),
        ("Payout", "Cuando Stripe transfiere el dinero de tus ventas a tu cuenta bancaria."),
        ("Microdepósito", "Transferencia pequeña ($1–$5 MXN) que Stripe hace para verificar propiedad de la cuenta."),
        ("Owner", "Persona con participación accionaria o de control significativo en la empresa (S.A. >25%, S. de R.L. >15%)."),
        ("Representante legal", "Persona autorizada para actuar en nombre de la empresa. La que firma contratos y opera cuentas."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        p.add_run(term + ". ").bold = True
        p.add_run(defn)

    # ----- Footer -----
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("¿Dudas? Captura de pantalla + mensaje a David. Si el problema es con datos fiscales o del SAT, esperar a tener la documentación correcta antes de avanzar.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    out = r"C:\Users\User\Documents\Click\docs\STRIPE_KYC_QLICK_FORMULARIO_AUTO_SELECCIONABLE.docx"
    doc.save(out)
    print(f"OK {out}")


if __name__ == "__main__":
    main()
