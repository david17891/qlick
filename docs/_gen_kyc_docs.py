#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera los 2 Word rellenables del KYC Stripe Qlick (Persona Física y Persona Moral)
con la nota del régimen RESICO ya incluida."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=RGBColor(0x00, 0x35, 0x54)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
    elif level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(11)
    return p


def add_warn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)


def add_ok(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x33, 0x77, 0x33)


def add_meta(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def make_table(doc, headers, rows, fill_header="E7F0F4"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], fill_header)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0x00, 0x35, 0x54)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
    return table


# ===== Document base: Persona Moral =====
def build_moral():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_heading(doc, "Stripe KYC — Persona Moral (S.A. o S. de R.L.)", level=0)
    add_meta(doc, "Qlick Marketing Digital · Última actualización: 2026-07-09 (incluye Régimen Simplificado de Confianza / RESICO)")

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Propósito: ").bold = True
    intro.add_run("Completar la verificación de identidad (KYC) de Stripe México para que Qlick pueda recibir pagos con tarjeta en producción, registrado como persona moral (empresa).")

    p = doc.add_paragraph()
    p.add_run("Para quién es esta guía: ").bold = True
    p.add_run("la persona que operará el dashboard de Stripe en nombre de David, con datos de la empresa.")

    p = doc.add_paragraph()
    p.add_run("Tiempo estimado: ").bold = True
    p.add_run("45 minutos a 1.5 horas si ya tiene todos los datos a mano; puede tomar más tiempo si debe pedir documentación al banco o contador.")

    p = doc.add_paragraph()
    p.add_run("Antes de empezar, confirme con David:").bold = True
    bullets = [
        "El tipo jurídico exacto (S.A. de C.V. o S. de R.L. de C.V.)",
        "Quiénes son los \"owners\" con más del 25% de participación (S.A.) o más del 15% (S. de R.L.) — van a tener que aparecer en el KYC",
        "Quién será el representante legal en el KYC (puede ser David o alguien de su confianza)",
        "El régimen fiscal declarado ante el SAT — Qlick está en Régimen Simplificado de Confianza (RESICO, código 626) según David. Verifique que el CSF diga \"Régimen Simplificado de Confianza\" antes de continuar.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    p = doc.add_paragraph()
    r = p.add_run("Si Qlick está registrada como Persona Física con Actividad Empresarial, use la otra guía.")
    r.italic = True

    add_heading(doc, "1. Datos que debe tener listos — bloque empresa", level=1)
    make_table(doc, ["#", "Dato", "Valor (llenar)", "Dónde obtenerlo"], [
        ["1", "Tipo jurídico exacto", "", "Acta constitutiva. Opciones válidas: S.A. de C.V. o S. de R.L. de C.V."],
        ["2", "Razón social completa", "", "Constancia de Situación Fiscal del SAT (CSF)"],
        ["3", "RFC de la empresa (12 caracteres)", "", "CSF del SAT"],
        ["4", "Domicilio fiscal completo", "", "CSF del SAT"],
        ["5", "Fecha de constitución", "", "Acta constitutiva"],
        ["5b", "Régimen fiscal del CSF", "626 - Régimen Simplificado de Confianza (RESICO)", "Confirmar en el CSF que diga RESICO con código 626"],
    ])

    add_ok(doc,
        "Sobre el régimen fiscal. El CSF del SAT incluye un campo \"Régimen fiscal\" con un código numérico. Para una persona moral en México los regímenes más comunes son: "
        "601 General de Ley Personas Morales (estándar), "
        "626 Régimen Simplificado de Confianza (RESICO) — es el que aplica a Qlick, "
        "603 Personas Morales con Fines no Lucrativos (ONGs, A.C.). "
        "Si su CSF trae otro código distinto a 626, confirme con David antes de avanzar.")

    add_heading(doc, "2. Datos del representante legal", level=1)
    make_table(doc, ["#", "Dato", "Valor (llenar)", "Dónde obtenerlo"], [
        ["6", "Nombre legal completo", "", "INE del representante (carácter por carácter)"],
        ["7", "RFC personal del representante (13 caracteres)", "", "CSF o INE"],
        ["8", "Fecha de nacimiento del representante", "", "INE"],
        ["9", "Dirección personal del representante", "", "Comprobante de domicilio (< 3 meses)"],
        ["10", "CURP del representante", "", "INE"],
    ])

    add_heading(doc, "3. Datos de los owners con participación significativa", level=1)
    add_warn(doc,
        "Quién califica como owner. S.A. de C.V.: cualquier persona con más del 25% de participación accionaria o con control financiero/operativo (CEO, CFO, Director). "
        "S. de R.L. de C.V.: cualquier socio con más del 15% de participación. "
        "Si Qlick tiene un solo dueño con 100%, es esa persona. Si tiene varios con menos del umbral, ninguno.")
    make_table(doc, ["#", "Dato por cada owner", "Valor (llenar)", "Tips"], [
        ["11", "Nombre completo", "", "Igual que en su INE"],
        ["12", "RFC personal (13 caracteres)", "", "Cada owner debe tener el suyo"],
        ["13", "Fecha de nacimiento", "", "dd/mm/aaaa"],
        ["14", "Dirección personal", "", "Calle + número + colonia + CP + municipio + estado"],
        ["15", "Porcentaje de participación", "", "Solo para personas con participación accionaria. Si es director sin acciones, dejar vacío"],
    ])

    add_heading(doc, "4. Datos bancarios", level=1)
    make_table(doc, ["#", "Dato", "Valor (llenar)", "Tips"], [
        ["16", "CLABE interbancaria de la empresa (18 dígitos)", "", "App del banco, \"datos para transferencia\""],
        ["17", "Titular de la cuenta", "", "Debe ser el nombre de la empresa, no una persona física"],
        ["18", "Nombre del banco", "", "BBVA, Banorte, Santander, HSBC, Scotiabank, Banregio. Evite neobanks"],
    ])

    add_heading(doc, "5. Datos del negocio y contacto público", level=1)
    make_table(doc, ["#", "Dato", "Valor (llenar)", "Sugerencia"], [
        ["19", "Statement descriptor (5–22 caracteres)", "", "QLICK.DIGITAL o QLICK CURSOS"],
        ["20", "URL del sitio web", "", "https://qlick.digital"],
        ["21", "Email de soporte", "", "hola@qlick.digital"],
        ["22", "Teléfono de soporte", "", "El que David designe"],
        ["23", "Categoría del negocio", "", "Education & Training"],
        ["24", "Descripción del negocio", "", "100–300 palabras"],
    ])

    p = doc.add_paragraph()
    p.add_run("Documentos que pueden pedir (preparar por si acaso):").bold = True
    for b in [
        "INE del representante (frente y vuelta, color, legible)",
        "Comprobante de domicilio del representante (< 3 meses)",
        "Constancia de Situación Fiscal del SAT (CSF)",
        "Acta constitutiva de la empresa",
        "Estado de cuenta bancario de la empresa",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "6. Pasos en el dashboard", level=1)

    steps = [
        ("Paso 1 — Iniciar sesión en Stripe",
         "URL: https://dashboard.stripe.com/login — Use las credenciales que David creó."),
        ("Paso 2 — Verificar que está en modo Live",
         "Arriba a la derecha debe decir \"Live data\". Si dice \"Test data\", cámbielo con el toggle."),
        ("Paso 3 — Iniciar el KYC",
         "URL: https://dashboard.stripe.com/account/onboarding — Click en \"Add information to start accepting live payments\"."),
        ("Paso 4 — Tipo de cuenta",
         "Elegir \"I'm setting up a business\"."),
        ("Paso 5 — País y tipo jurídico",
         "País: Mexico. Tipo jurídico: S.A. de C.V. o S. de R.L. de C.V. (el que aplique)."),
        ("Paso 6 — Datos de la empresa",
         "Llenar con datos #1 a #5b de la sección 1. Confirmar que el régimen fiscal sea 626 si Qlick está en RESICO. El nombre legal debe coincidir carácter por carácter con el CSF. \"S. de R.L.\" con punto difiere de \"S de RL\" sin él."),
        ("Paso 7 — Datos del negocio",
         "Nombre legal: la razón social. Categoría: Education & Training. "
         "Descripción sugerida: \"Plataforma mexicana de educación online en marketing digital. Ofrece cursos grabados, workshops en vivo y material descargable para emprendedores y profesionales. Mercado: México y Latinoamérica hispanohablante. Sitio web: qlick.digital.\" "
         "Sitio web: https://qlick.digital. Support email y phone: los definidos por David."),
        ("Paso 8 — Datos del representante legal",
         "Llenar con datos #6 a #10 de la sección 2. El representante debe ser la persona autorizada para operar la cuenta."),
        ("Paso 9 — Owners adicionales (si aplica)",
         "Stripe le preguntará si hay otros owners con participación significativa. Si aplica los umbrales de la sección 3, agregue uno por uno. Datos por owner: #11 a #15."),
        ("Paso 10 — Statement descriptor",
         "URL: https://dashboard.stripe.com/settings/public — Escribir el descriptor (#19)."),
        ("Paso 11 — Cuenta bancaria de la empresa",
         "País: Mexico. Moneda: MXN. Account number: CLABE (#16, 18 dígitos). Account holder: nombre de la empresa (#17)."),
        ("Paso 12 — Subir documentos (si los pide)",
         "Si Stripe pide verificación adicional: INE del representante (foto, frente y vuelta), CSF del SAT (PDF o foto legible), estado de cuenta bancario de la empresa (con CLABE visible). Foto con buena luz, el documento plano sobre mesa."),
        ("Paso 13 — Esperar activación",
         "URL: https://dashboard.stripe.com/account — Sección \"Account status\". 🟡 Restricted: falta algo. 🟢 Complete: listo."),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(body)

    add_heading(doc, "7. Gotchas importantes", level=1)
    add_warn(doc,
        "Razón social carácter por carácter. \"S. de R.L.\" vs \"S de RL\" puede parecer lo mismo pero Stripe rechaza si no coincide con el CSF.")
    add_warn(doc,
        "RESICO (régimen 626). Qlick está dado de alta en el Régimen Simplificado de Confianza. Verifique que el CSF diga \"Régimen Simplificado de Confianza\" con código 626. "
        "Si Stripe le pide el régimen fiscal y usted pone otro código (601 General de Ley, etc), la verificación se complica. Si su CSF dice otro régimen, frene y avísele a David.")
    add_warn(doc,
        "Owners. Si Stripe le pregunta por otros owners y Qlick tiene un solo dueño que es David, indique \"No\" o agregue a David como owner con 100% de participación.")
    add_warn(doc,
        "Titular de la cuenta bancaria. Debe ser la empresa (razón social), no una persona. Si está aperturada a nombre de una persona, hay que cambiarla antes de continuar.")
    add_warn(doc, "CLABE ≠ número de tarjeta. 18 dígitos, no 16. Verifique que empieza con código de banco (012, 072, etc).")
    add_warn(doc, "Microdepósitos. Stripe transfiere $1–$5 MXN a la cuenta de la empresa para verificar propiedad. Tarda 1–3 días hábiles. El proceso se pausa hasta confirmarlos.")
    add_warn(doc, "No cambiable después. País, tipo jurídico, RFC de empresa y razón social no se pueden editar una vez confirmados. Verifique tres veces antes de avanzar.")

    add_heading(doc, "8. Qué le aviso a David cuando termine", level=1)
    p = doc.add_paragraph("Cuando el estado esté 🟢 Complete, envíele a David:")
    for b in [
        "Screenshot del estado verde en https://dashboard.stripe.com/account",
        "Confirmación de que puede generar keys live en https://dashboard.stripe.com/apikeys (toggle \"Live data\")",
        "Confirmación de métodos de pago activos: Cards (obligatorio), OXXO (opcional pero recomendado), SPEI / customer balance (opcional)",
        "URL del webhook endpoint que David va a registrar: https://www.qlick.digital/api/webhooks/stripe — eventos requeridos: checkout.session.completed, checkout.session.async_payment_succeeded, checkout.session.async_payment_failed, checkout.session.expired, charge.refunded",
    ]:
        doc.add_paragraph(b, style="List Number")

    add_heading(doc, "9. Glosario mínimo", level=1)
    glossary = [
        ("KYC", "\"Know Your Customer\" — verificación obligatoria de identidad por ley."),
        ("RFC", "Registro Federal de Contribuyentes. Personas físicas = 13 caracteres. Empresas = 12."),
        ("CSF", "Constancia de Situación Fiscal del SAT. Documento oficial con RFC, razón social, domicilio fiscal, régimen fiscal."),
        ("S.A. de C.V.", "Sociedad Anónima de Capital Variable. Tipo jurídico común para empresas en México."),
        ("S. de R.L. de C.V.", "Sociedad de Responsabilidad Limitada de Capital Variable."),
        ("Owner (accionista/socio)", "Persona con participación accionaria o de control significativo en la empresa."),
        ("Representante legal", "Persona autorizada para actuar en nombre de la empresa. La que firma contratos y opera cuentas."),
        ("CLABE", "Clave Bancaria Estandarizada. 18 dígitos. La usa Stripe para depositar las ventas."),
        ("Statement descriptor", "Texto corto que aparece en el estado de cuenta del cliente."),
        ("Microdepósito", "Transferencia pequeña ($1–$5 MXN) que Stripe hace para verificar propiedad de la cuenta."),
        ("RESICO", "Régimen Simplificado de Confianza. Régimen fiscal opcional del SAT para personas físicas y morales con ingresos menores a ciertos topes. Códigos: 626 (persona moral) o 625 (persona física con actividad empresarial). Qlick está en RESICO."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        p.add_run(term + ". ").bold = True
        p.add_run(defn)

    out = r"C:\Users\User\Documents\Click\docs\STRIPE_KYC_QLICK_PERSONA_MORAL.docx"
    doc.save(out)
    print(f"OK {out}")


# ===== Document base: Persona Física con AE =====
def build_fisica():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_heading(doc, "Stripe KYC — Persona Física con Actividad Empresarial", level=0)
    add_meta(doc, "Qlick Marketing Digital · Última actualización: 2026-07-09 (incluye Régimen Simplificado de Confianza / RESICO)")

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Propósito: ").bold = True
    intro.add_run("Completar la verificación de identidad (KYC) de Stripe México para que Qlick pueda recibir pagos con tarjeta en producción.")

    p = doc.add_paragraph()
    p.add_run("Para quién es esta guía: ").bold = True
    p.add_run("la persona que operará el dashboard de Stripe en nombre de David.")

    p = doc.add_paragraph()
    p.add_run("Tiempo estimado: ").bold = True
    p.add_run("30–45 minutos si ya tiene todos los datos a mano; 1–2 horas si debe ir a buscarlos al SAT o al banco.")

    p = doc.add_paragraph()
    p.add_run("Antes de empezar, confirme con David:").bold = True
    for b in [
        "Que el tipo jurídico correcto es Persona Física con Actividad Empresarial. Si Qlick ya está constituida como S.A. o S. de R.L., use la otra guía.",
        "El régimen fiscal declarado ante el SAT — David indica que Qlick está en Régimen Simplificado de Confianza (RESICO, código 625) para persona física con actividad empresarial. Verifique que su CSF diga \"Régimen Simplificado de Confianza\" antes de continuar.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "1. Datos que debe tener listos", level=1)
    make_table(doc, ["#", "Dato", "Valor (llenar)", "Dónde obtenerlo"], [
        ["1", "Tu nombre legal completo", "", "Tu INE — carácter por carácter, sin abreviaturas"],
        ["2", "Tu RFC personal (13 caracteres)", "", "Constancia de Situación Fiscal del SAT, o tu INE si lo trae"],
        ["3", "Tu fecha de nacimiento", "", "INE"],
        ["4", "Tu dirección personal", "", "Comprobante de domicilio (< 3 meses: estado de cuenta, CFE, Telmex)"],
        ["5", "CLABE interbancaria (18 dígitos)", "", "App de tu banco, sección \"datos para transferencia\" o \"CLABE\""],
        ["6", "Titular de la cuenta bancaria", "", "Estado de cuenta bancario o app del banco"],
        ["7", "Nombre del banco", "", "BBVA, Banorte, Santander, HSBC, Scotiabank, Banregio. Evite neobanks (Nu, HeyBanco, Spin, Mercado Pago wallet)"],
        ["8", "Statement descriptor (5–22 caracteres)", "", "Decisión del equipo. Sugerencia: QLICK.DIGITAL"],
        ["9", "URL del sitio web", "", "https://qlick.digital"],
        ["10", "Email de soporte", "", "hola@qlick.digital (o el que David indique)"],
        ["11", "Teléfono de soporte", "", "El que David designe"],
        ["11b", "Régimen fiscal del CSF", "625 - Régimen Simplificado de Confianza (Persona Física)", "Confirmar en el CSF que diga RESICO con código 625"],
    ])

    add_ok(doc,
        "Sobre el régimen fiscal. El CSF del SAT incluye un campo \"Régimen fiscal\" con un código numérico. Para una persona física con actividad empresarial los códigos más comunes son: "
        "612 Persona Física con Actividad Empresarial (estándar), "
        "625 Régimen Simplificado de Confianza (RESICO) — es el que aplica a Qlick. "
        "Si su CSF trae otro código distinto a 625, confirme con David antes de avanzar.")

    p = doc.add_paragraph()
    p.add_run("Documentos que pueden pedirle (preparar por si acaso):").bold = True
    for b in [
        "Foto de tu INE frente y vuelta (color, legible)",
        "Comprobante de domicilio personal (< 3 meses)",
        "Constancia de Situación Fiscal del SAT (CSF) — solo si Stripe te pide verificar empresa",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "2. Pasos en el dashboard", level=1)
    steps = [
        ("Paso 1 — Iniciar sesión en Stripe",
         "URL: https://dashboard.stripe.com/login — Use las credenciales que David creó."),
        ("Paso 2 — Verificar que está en modo Live",
         "Arriba a la derecha debe decir \"Live data\". Si dice \"Test data\", cámbielo con el toggle. Stripe en modo test no sirve para producción real."),
        ("Paso 3 — Iniciar el KYC",
         "URL: https://dashboard.stripe.com/account/onboarding — Click en \"Add information to start accepting live payments\"."),
        ("Paso 4 — Tipo de cuenta",
         "Elegir \"I'm setting up a business\" (empresa, no individual). No elija \"individual\" porque después tendrá que reiniciar el proceso."),
        ("Paso 5 — País y tipo jurídico",
         "País: Mexico. Tipo jurídico: Persona Física con Actividad Empresarial."),
        ("Paso 6 — Datos personales",
         "Llenar el formulario con tus datos (los puntos 1 a 4 de la tabla de arriba)."),
        ("Paso 7 — Datos del negocio",
         "En la sección \"Business profile\". Nombre legal: tu nombre completo. Categoría: Education & Training. "
         "Descripción sugerida: \"Plataforma mexicana de educación online en marketing digital. Ofrece cursos grabados, workshops en vivo y material descargable para emprendedores y profesionales. Mercado: México y Latinoamérica hispanohablante. Sitio web: qlick.digital.\" "
         "Sitio web: https://qlick.digital. Support email y phone: el que David indique."),
        ("Paso 8 — Statement descriptor",
         "URL: https://dashboard.stripe.com/settings/public — Escribir el descriptor elegido (punto 8). Aparece en el estado de cuenta del cliente."),
        ("Paso 9 — Cuenta bancaria",
         "En la sección de payouts. País: Mexico. Moneda: MXN. Account number: CLABE (18 dígitos). Account holder: titular de la cuenta."),
        ("Paso 10 — Subir documentos (si los pide)",
         "Stripe valida automáticamente la mayoría de los datos. Si no puede, le pide: INE vigente (foto, frente y vuelta), comprobante de domicilio. Súbalos desde su celular, con buena luz, legibles."),
        ("Paso 11 — Esperar activación",
         "URL: https://dashboard.stripe.com/account — Sección \"Account status\". 🟡 Restricted: falta algo. 🟢 Complete: listo. Avise a David."),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(body)

    add_heading(doc, "3. Gotchas importantes", level=1)
    add_warn(doc, "CLABE ≠ número de tarjeta. La CLABE son 18 dígitos y empieza con el código del banco (012 = BBVA, 072 = Banorte, 014 = Santander). El número de tarjeta son 16 dígitos y NO sirve. Si Stripe lo rechaza, vuelva a verificar.")
    add_warn(doc, "No usar neobanks. Nu, HeyBanco, Spin y Mercado Pago wallet a veces no funcionan con Stripe payouts. Si su cuenta principal es en uno de esos, abra una cuenta en BBVA o Banorte tradicional.")
    add_warn(doc, "El nombre debe coincidir. El nombre que pongas en Stripe debe coincidir carácter por carácter con tu INE. Variaciones como \"Pérez\" vs \"Perez\" causan rechazo.")
    add_warn(doc, "Una vez enviado, no se puede cambiar. El país, el nombre legal y el RFC no son editables después. Verifique tres veces antes de avanzar.")
    add_warn(doc, "Microdepósitos. Stripe hace 2 transferencias de $1–$5 MXN a tu cuenta para verificar que es tuya. Tarda 1–3 días hábiles. Hasta que las confirmen, el proceso queda pausado.")
    add_warn(doc, "RESICO (régimen 625). Qlick está dado de alta en el Régimen Simplificado de Confianza. Verifique que el CSF diga \"Régimen Simplificado de Confianza\" con código 625. Si su CSF dice otro régimen, frene y avísele a David.")

    add_heading(doc, "4. Qué le aviso a David cuando termine", level=1)
    p = doc.add_paragraph("Cuando el estado esté 🟢 Complete, envíele a David:")
    for b in [
        "Screenshot del estado verde en https://dashboard.stripe.com/account",
        "Confirmación de que puede generar keys live en https://dashboard.stripe.com/apikeys (toggle \"Live data\")",
        "Confirmación de que los métodos de pago están activos: Cards (obligatorio), OXXO (opcional pero recomendado), SPEI / customer balance (opcional)",
        "URL donde está el webhook endpoint que David va a registrar: https://www.qlick.digital/api/webhooks/stripe — eventos: checkout.session.completed, checkout.session.async_payment_succeeded, checkout.session.async_payment_failed, checkout.session.expired, charge.refunded",
    ]:
        doc.add_paragraph(b, style="List Number")

    add_heading(doc, "5. Glosario mínimo", level=1)
    glossary = [
        ("KYC", "\"Know Your Customer\" — verificación obligatoria de identidad por ley."),
        ("RFC", "Registro Federal de Contribuyentes (México). Personas físicas = 13 caracteres. Empresas = 12."),
        ("CLABE", "Clave Bancaria Estandarizada. 18 dígitos. La usa Stripe para depositarte el dinero de las ventas."),
        ("Statement descriptor", "Texto corto que aparece en el estado de cuenta del cliente cuando le cobrás."),
        ("Payout", "Cuando Stripe transfiere el dinero de tus ventas a tu cuenta bancaria."),
        ("Microdepósito", "Transferencia pequeña ($1–$5 MXN) que Stripe hace para verificar propiedad de la cuenta."),
        ("RESICO", "Régimen Simplificado de Confianza. Régimen fiscal opcional del SAT para personas físicas y morales con ingresos menores a ciertos topes. Códigos: 626 (persona moral) o 625 (persona física con actividad empresarial). Qlick está en RESICO."),
    ]
    for term, defn in glossary:
        p = doc.add_paragraph()
        p.add_run(term + ". ").bold = True
        p.add_run(defn)

    out = r"C:\Users\User\Documents\Click\docs\STRIPE_KYC_QLICK_PERSONA_FISICA.docx"
    doc.save(out)
    print(f"OK {out}")


if __name__ == "__main__":
    build_moral()
    build_fisica()
