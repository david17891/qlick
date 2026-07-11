#!/usr/bin/env python
"""
Genera 2 archivos .docx con tablas rellenables para el KYC de Stripe Qlick.
- PERSONA_FISICA: para Persona Física con Actividad Empresarial.
- PERSONA_MORAL:   para S.A. de C.V. o S. de R.L. de C.V.

Campos en una columna "Dato" (label) y otra "Valor" (vacía) para que
David (o el operador) escriba. Estructura formal, español mexicano neutro.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path("docs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Paleta neutra Qlick (azul oscuro corporativo)
NAVY = RGBColor(0x00, 0x35, 0x54)
LIGHT_BLUE = "E7F0F4"
SECTION_PAD = (4, 6, 4, 6)  # top, left, bottom, right (twips approx)


def set_cell_background(cell, color_hex: str):
    """Aplica color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """Aplica bordes grises finos a la celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_title(doc, text, size=22, color=NAVY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_subtitle(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_section(doc, text, size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_paragraph_text(doc, text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def add_bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(size)


def add_data_table(doc, headers, rows):
    """Tabla con header (label + valor)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr[i], "003554")
        set_cell_borders(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xA3, 0x33, 0x33)
    run.font.bold = True
    return p


def doc_base():
    doc = Document()
    # Establecer márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    return doc


# ============================================================
#  PERSONA FÍSICA CON ACTIVIDAD EMPRESARIAL
# ============================================================
def build_persona_fisica():
    doc = doc_base()
    add_title(doc, "Stripe KYC — Persona Física con Actividad Empresarial")
    add_subtitle(doc, "Qlick Marketing Digital — Formulario para que David (o el operador) complete antes de enviar a Stripe.")

    # Intro
    add_paragraph_text(
        doc,
        "Esta guía agrupa todos los datos que Stripe va a pedir para verificar a Qlick "
        "como Persona Física con Actividad Empresarial. Antes de empezar, confirme con David "
        "que el tipo jurídico correcto sea Persona Física. Si Qlick está constituida como S.A. "
        "o S. de R.L., use la otra guía.",
        bold=False,
    )

    add_section(doc, "1. Datos personales del titular / representante")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["Nombre legal completo (como aparece en tu INE)", ""],
            ["RFC personal (13 caracteres)", ""],
            ["Fecha de nacimiento (dd/mm/aaaa)", ""],
            ["CURP", ""],
            ["Dirección personal — Calle y número", ""],
            ["Dirección personal — Colonia", ""],
            ["Dirección personal — CP", ""],
            ["Dirección personal — Municipio / Alcaldía", ""],
            ["Dirección personal — Estado", ""],
            ["Teléfono personal", ""],
            ["Email personal", ""],
        ],
    )

    add_section(doc, "2. Datos bancarios (cuenta a la que Stripe te deposita)")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["CLABE interbancaria (18 dígitos, NO el número de tarjeta)", ""],
            ["Titular de la cuenta bancaria", ""],
            ["Nombre del banco", ""],
            ["Tipo de cuenta (ahorro / nómina / etc.)", ""],
        ],
    )

    add_warning(doc, "⚠ La CLABE son 18 dígitos. NO confundas con tu número de tarjeta (16 dígitos).")

    add_section(doc, "3. Datos del negocio")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["Razón social / nombre del negocio (puede ser tu nombre)", ""],
            ["Categoría sugerida", "Education & Training"],
            ["Descripción del negocio (100–300 palabras)", ""],
            ["URL del sitio web", "https://qlick.digital"],
            ["Email de soporte (aparece en recibos al cliente)", ""],
            ["Teléfono de soporte", ""],
            ["Statement descriptor (5–22 caracteres, sugerencia QLICK.DIGITAL)", ""],
        ],
    )

    add_paragraph_text(
        doc,
        "Plantilla sugerida para la descripción: «Plataforma mexicana de educación online en marketing "
        "digital. Ofrece cursos grabados, workshops en vivo y material descargable para emprendedores "
        "y profesionales. Mercado: México y Latinoamérica hispanohablante. Sitio web: qlick.digital.»",
        bold=False,
    )

    add_section(doc, "4. Documentos que pueden pedirte (preparar por si acaso)")
    add_bullets(doc, [
        "Foto de tu INE vigente, frente y vuelta (color, legible, plano sobre mesa).",
        "Comprobante de domicilio personal, menor a 3 meses (recibo de CFE, Telmex, estado de cuenta bancario).",
        "Constancia de Situación Fiscal del SAT (CSF), si Stripe la pide para validar empresa.",
    ])

    add_section(doc, "5. Pasos en el dashboard de Stripe")
    add_bullets(doc, [
        "Abre https://dashboard.stripe.com/login e inicia sesión con las credenciales que David te pasó.",
        "Verifica que arriba a la derecha diga «Live data». Si dice «Test data», cámbialo con el toggle.",
        "Ve a https://dashboard.stripe.com/account/onboarding y haz clic en «Add information to start accepting live payments».",
        "Selecciona «I'm setting up a business» (empresa, no individual).",
        "País: Mexico. Tipo jurídico: Persona Física con Actividad Empresarial.",
        "Llena el formulario con los datos de las secciones 1, 2 y 3 de este documento.",
        "En la sección de Business profile, la categoría debe ser «Education & Training».",
        "Confirma el statement descriptor en https://dashboard.stripe.com/settings/public.",
        "Agrega la cuenta bancaria con la CLABE.",
        "Si Stripe te pide verificación, sube INE y comprobante de domicilio.",
        "Espera a que el estado en https://dashboard.stripe.com/account pase a verde (Complete).",
        "Avísale a David para que genere las claves de producción y registre el webhook.",
    ])

    add_section(doc, "6. Gotchas a evitar")
    add_bullets(doc, [
        "CLABE de 18 dígitos ≠ número de tarjeta de 16 dígitos. Verifica antes de enviar.",
        "Cuenta bancaria en banco tradicional (BBVA, Banorte, Santander). No uses neobanks.",
        "El nombre legal debe coincidir carácter por carácter con tu INE.",
        "Una vez confirmados, el país, el tipo jurídico y el RFC no se pueden cambiar.",
        "Stripe hace microdepósitos de $1–$5 MXN para verificar la cuenta. Tarda 1–3 días hábiles.",
    ])

    add_section(doc, "7. Datos finales para avisarle a David")
    add_data_table(
        doc,
        headers=["Confirmación a enviar", "Valor"],
        rows=[
            ["Estado del KYC al terminar (verde / amarillo / rojo)", ""],
            ["¿Microdepósitos confirmados? (sí / no / pendiente)", ""],
            ["Últimos 4 dígitos de la CLABE agregada", ""],
            ["Statement descriptor configurado", ""],
            ["Email de soporte registrado", ""],
            ["¿Listo para que David genere las claves live? (sí / no)", ""],
        ],
    )

    output = OUTPUT_DIR / "STRIPE_KYC_QLICK_PERSONA_FISICA.docx"
    doc.save(output)
    print(f"OK: {output}")


# ============================================================
#  PERSONA MORAL (S.A. o S. de R.L.)
# ============================================================
def build_persona_moral():
    doc = doc_base()
    add_title(doc, "Stripe KYC — Persona Moral (S.A. o S. de R.L.)")
    add_subtitle(doc, "Qlick Marketing Digital — Formulario para que David (o el operador) complete antes de enviar a Stripe.")

    add_paragraph_text(
        doc,
        "Esta guía agrupa todos los datos que Stripe va a pedir para verificar a Qlick "
        "como persona moral. Antes de empezar, confirme con David: (a) el tipo jurídico exacto "
        "(S.A. de C.V. o S. de R.L. de C.V.), (b) quiénes son los owners con más del 25% (S.A.) o "
        "más del 15% (S. de R.L.) de participación, y (c) quién será el representante legal.",
        bold=False,
    )

    add_section(doc, "1. Datos de la empresa")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["Tipo jurídico exacto (S.A. de C.V. o S. de R.L. de C.V.)", ""],
            ["Razón social completa (carácter por carácter, del CSF)", ""],
            ["RFC de la empresa (12 caracteres, del CSF)", ""],
            ["Domicilio fiscal — Calle y número (del CSF)", ""],
            ["Domicilio fiscal — Colonia (del CSF)", ""],
            ["Domicilio fiscal — CP (del CSF)", ""],
            ["Domicilio fiscal — Municipio (del CSF)", ""],
            ["Domicilio fiscal — Estado (del CSF)", ""],
            ["Fecha de constitución (dd/mm/aaaa, del acta constitutiva)", ""],
        ],
    )
    add_warning(doc, "⚠ La razón social debe coincidir carácter por carácter con el CSF. «S. de R.L.» con punto difiere de «S de RL» sin él.")

    add_section(doc, "2. Datos del representante legal")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["Nombre legal completo del representante", ""],
            ["RFC personal del representante (13 caracteres)", ""],
            ["Fecha de nacimiento del representante (dd/mm/aaaa)", ""],
            ["CURP del representante", ""],
            ["Dirección personal del representante — Calle y número", ""],
            ["Dirección personal del representante — Colonia", ""],
            ["Dirección personal del representante — CP", ""],
            ["Dirección personal del representante — Municipio", ""],
            ["Dirección personal del representante — Estado", ""],
            ["Teléfono del representante", ""],
            ["Email del representante", ""],
            ["Cargo (Director / CEO / Representante Legal)", ""],
        ],
    )

    add_section(doc, "3. Datos de los owners con participación significativa")
    add_paragraph_text(
        doc,
        "Aplica solo para personas con más del 25% (S.A.) o 15% (S. de R.L.) de participación, o con "
        "control financiero de la empresa. Si solo hay un dueño con 100%, es esa persona. Si no "
        "alcanza el umbral nadie, dejar vacío.",
        bold=False,
    )
    add_data_table(
        doc,
        headers=["Dato", "Owner 1", "Owner 2 (si aplica)"],
        rows=[
            ["Nombre completo", "", ""],
            ["RFC personal (13 caracteres)", "", ""],
            ["Fecha de nacimiento (dd/mm/aaaa)", "", ""],
            ["Dirección personal", "", ""],
            ["Porcentaje de participación", "", ""],
        ],
    )

    add_section(doc, "4. Datos bancarios de la empresa")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["CLABE interbancaria de la empresa (18 dígitos)", ""],
            ["Titular de la cuenta (debe ser la razón social, no una persona)", ""],
            ["Nombre del banco", ""],
            ["Tipo de cuenta", ""],
        ],
    )
    add_warning(doc, "⚠ Si la cuenta bancaria está a nombre de una persona física, hay que cambiarla antes de continuar.")

    add_section(doc, "5. Datos del negocio y contacto público")
    add_data_table(
        doc,
        headers=["Dato", "Valor"],
        rows=[
            ["Categoría", "Education & Training"],
            ["Descripción del negocio (100–300 palabras)", ""],
            ["URL del sitio web", "https://qlick.digital"],
            ["Email de soporte (aparece en recibos)", ""],
            ["Teléfono de soporte", ""],
            ["Statement descriptor (sugerencia QLICK.DIGITAL)", ""],
        ],
    )

    add_paragraph_text(
        doc,
        "Plantilla sugerida: «Plataforma mexicana de educación online en marketing digital. Ofrece "
        "cursos grabados, workshops en vivo y material descargable para emprendedores y profesionales. "
        "Mercado: México y Latinoamérica hispanohablante. Sitio web: qlick.digital.»",
        bold=False,
    )

    add_section(doc, "6. Documentos que pueden pedirte (preparar por si acaso)")
    add_bullets(doc, [
        "INE del representante legal, frente y vuelta (color, legible).",
        "Comprobante de domicilio del representante, menor a 3 meses.",
        "Constancia de Situación Fiscal del SAT (CSF) de la empresa.",
        "Acta constitutiva de la empresa.",
        "Estado de cuenta bancario de la empresa (para verificar CLABE).",
    ])

    add_section(doc, "7. Pasos en el dashboard de Stripe")
    add_bullets(doc, [
        "Abre https://dashboard.stripe.com/login.",
        "Verifica que arriba a la derecha diga «Live data» (no «Test data»).",
        "Ve a https://dashboard.stripe.com/account/onboarding y haz clic en «Add information to start accepting live payments».",
        "Selecciona «I'm setting up a business».",
        "País: Mexico. Tipo jurídico: S.A. de C.V. o S. de R.L. de C.V.",
        "Llena los datos de la empresa (sección 1 de este documento).",
        "Llena los datos del representante legal (sección 2).",
        "Si Stripe lo pide, agrega a cada owner (sección 3) uno por uno.",
        "En Business profile, categoría: «Education & Training».",
        "Confirma el statement descriptor en https://dashboard.stripe.com/settings/public.",
        "Agrega la cuenta bancaria de la empresa con la CLABE.",
        "Si Stripe te pide verificación, sube los documentos de la sección 6.",
        "Espera a que el estado en https://dashboard.stripe.com/account pase a verde (Complete).",
        "Avísale a David para que genere las claves y registre el webhook.",
    ])

    add_section(doc, "8. Gotchas a evitar")
    add_bullets(doc, [
        "Razón social carácter por carácter. Variaciones como «S. de R.L.» vs «S de RL» causan rechazo.",
        "El titular de la cuenta bancaria debe ser la empresa, no una persona.",
        "CLABE de 18 dígitos, no número de tarjeta.",
        "Stripe hace microdepósitos de $1–$5 MXN, tarda 1–3 días hábiles.",
        "País, tipo jurídico, RFC de empresa y razón social no se pueden cambiar después.",
        "Si Qlick tiene un solo dueño, es ese owner con 100% — no lo omitas.",
    ])

    add_section(doc, "9. Datos finales para avisarle a David")
    add_data_table(
        doc,
        headers=["Confirmación a enviar", "Valor"],
        rows=[
            ["Estado del KYC al terminar (verde / amarillo / rojo)", ""],
            ["¿Microdepósitos confirmados? (sí / no / pendiente)", ""],
            ["Últimos 4 dígitos de la CLABE de la empresa", ""],
            ["Statement descriptor configurado", ""],
            ["Email de soporte registrado", ""],
            ["¿Listo para que David genere las claves live? (sí / no)", ""],
        ],
    )

    output = OUTPUT_DIR / "STRIPE_KYC_QLICK_PERSONA_MORAL.docx"
    doc.save(output)
    print(f"OK: {output}")


if __name__ == "__main__":
    build_persona_fisica()
    build_persona_moral()
